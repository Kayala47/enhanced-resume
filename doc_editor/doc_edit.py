import argparse
from typing import List, Tuple
from docx import Document

SORTED_SKILLS_TEST = ["Python", "React.js", "Swift", "C++"]


def sort_skills(job_skills, resume_skills):
    for js in reversed(job_skills):
        for i, rs in enumerate(resume_skills):
            if js.upper() == rs.upper():
                resume_skills.insert(0, resume_skills.pop(i))

    return resume_skills


def coallesece_skils(doc, idxs, text_lists):

    for i, list in zip(idxs, text_lists):
        print(list)
        doc.paragraphs[i].text = ", ".join(list)

def sort_exp(job_exp, resume_exp):
    sortedList = []
    for paragraph in resume_exp:
        jobdict = dict(paragraph, [0]*len(paragraph))
        for i, jex in enumerate(reversed(job_exp)):
            for rex in paragraph:
                if jex in rex:
                    jobdict[rex] += i
        jobdict = sorted(jobdict)
        sortedList += [list(jobdict.keys())]
    return sortedList
        


def scrape_doc(
    doc,
) -> Tuple[List[str], List[int], List[str], List[int], List[str], List[int]]:

    exp = []
    exp_idx = []

    attr = []
    attr_idx = []

    styles = doc.styles

    skills_lists = []
    skills_idxs = []
    last_style = "Heading 1"
    stage = 0  # whether we're scraping attr, exp, or skills

    for i, p in enumerate(doc.paragraphs):
        if p.style == styles["Heading 2"]:
            # time for a context switch!
            # TODO add exp and attr

            if p.text == "SKILLS":
                stage = 3
                # but don't scrape this line
                continue

            else:
                stage = 4

        if stage == 3:
            print("in stage 3")
            # scraping skills
            if p.style != styles["Heading 3"]:
                # don't want to collect these, but collect the actual info for sure
                skills_lists.append(p.text.split(", "))
                skills_idxs.append(i)

    return skills_lists, skills_idxs, exp, exp_idx, attr, attr_idx


def main(args):
    # sorted_skills, sorted_exp, sorted_attr = get_sorted_lists(args.search)
    doc = Document(args.resume_filename)

    (
        present_skills,
        skills_idx,
        present_exp,
        exp_idx,
        present_attr,
        attr_idx,
    ) = scrape_doc(doc)

    new_skills = []
    for skill_list in present_skills:
        new_skills.append(sort_skills(SORTED_SKILLS_TEST, skill_list))

    print(new_skills)

    coallesece_skils(doc, skills_idx, new_skills)

    doc.save("output_resume.docx")

    # new_exp = sort_experiences(sorted_exp, present_exp)

    # write_paragraphs(new_skills, skills_idx, args.filename)
    # write_paragraphs(new_exp, exp_idx, args.filename)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "search",
        type=str,
        help="Job URL or job title that this resume is to be tailored for",
    )
    parser.add_argument("resume_filename", type=str, help="Path to current resume")
    args = parser.parse_args()
    main(args)
